#!/usr/bin/env python3
"""
Genera il template DOCX del giustificativo (attestazione di presenza)
partendo dal template della fattura, di cui riusa la carta intestata:
tutto ciò che precede il blocco destinatario/corpo fattura viene mantenuto
(logo e intestazione compresi), il resto viene sostituito dal corpo
dell'attestazione con i placeholder docxtpl (riempiti a runtime da
GET /api/clients/<id>/giustificativo in clienti_api.py).

Uso:
    # Template built-in (base: backend/app/templates/invoice_template.docx)
    python3 scripts/generate_giustificativo_template.py

    # Template custom di produzione (base con logo e dati già impostati)
    python3 scripts/generate_giustificativo_template.py \
        --base custom_template/invoice_template.docx \
        --out custom_template/giustificativo_template.docx

Richiede python-docx (già presente nei requirements del backend).
"""
import argparse
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BLU_SCURO = RGBColor(0x1F, 0x3A, 0x5F)
GRIGIO = RGBColor(0x59, 0x59, 0x59)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DEFAULT_BASE = os.path.join(REPO_ROOT, 'backend', 'app', 'templates', 'invoice_template.docx')
DEFAULT_OUT = os.path.join(REPO_ROOT, 'backend', 'app', 'templates', 'giustificativo_template.docx')

# Il primo elemento del body che contiene uno di questi marker segna l'inizio
# della parte "fattura" da eliminare; tutto ciò che precede è carta intestata.
MARKERS = ('cliente_nome', 'FATTURA N.')


def strip_invoice_body(doc):
    """Rimuove dal body tutto a partire dal blocco destinatario/fattura."""
    body = doc.element.body
    removing = False
    for element in list(body):
        if element.tag == qn('w:sectPr'):
            continue
        if not removing:
            text = ''.join(element.itertext())
            if any(marker in text for marker in MARKERS):
                removing = True
        if removing:
            body.remove(element)
    if not removing:
        raise SystemExit(
            "Blocco fattura non trovato nel template base: nessun elemento "
            f"contiene i marker {MARKERS}. Il file è un template fattura valido?")


def para(doc, text='', size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
         color=None, space_after=2, space_before=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    return p


def add_run(p, text, size=12, bold=False):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return run


def append_giustificativo_body(doc):
    # --- Titolo ---
    titolo = para(doc, size=14, space_before=26, space_after=22,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    run = add_run(titolo, 'ATTESTAZIONE DI PRESENZA', size=14, bold=True)
    run.font.color.rgb = BLU_SCURO
    # Leggera spaziatura tra le lettere
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), '30')
    rPr.append(spacing)

    # --- Corpo ---
    corpo = doc.add_paragraph()
    corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    corpo.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    corpo.paragraph_format.space_after = Pt(14)
    add_run(corpo, '{{ articolo_titolo }} ')
    add_run(corpo, '{{ cliente_nome }} {{ cliente_cognome }}', bold=True)
    add_run(corpo, ', {{ nato_nata }} a {{ cliente_luogo_nascita }} il '
                   '{{ cliente_data_nascita }}, residente a {{ cliente_residenza }}, '
                   'codice fiscale {{ cliente_codice_fiscale }}, si è '
                   '{{ presentato_presentata }} presso il mio studio in data ')
    add_run(corpo, '{{ data_prestazione }}', bold=True)
    add_run(corpo, ', dalle ore ')
    add_run(corpo, '{{ ora_inizio }}', bold=True)
    add_run(corpo, ' alle ore ')
    add_run(corpo, '{{ ora_fine }}', bold=True)
    add_run(corpo, ', per sottoporsi a prestazione sanitaria.')

    rilascio = doc.add_paragraph()
    rilascio.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rilascio.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    rilascio.paragraph_format.space_after = Pt(14)
    add_run(rilascio, "Si rilascia la presente attestazione su richiesta "
                      "dell'{{ interessato_interessata }}, su carta semplice e in "
                      "copia unica, per gli usi consentiti dalla legge.")

    para(doc, '{{ luogo_data_rilascio }}', size=12, space_before=18, space_after=0)

    # --- Firma ---
    para(doc, 'Timbro e firma', size=12, align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_before=24, space_after=0)
    para(doc, '{{ intestatario_titolo }}', size=12, italic=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, color=GRIGIO, space_before=6, space_after=48)
    para(doc, '_______________________________', size=12,
         align=WD_ALIGN_PARAGRAPH.RIGHT)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('--base', default=DEFAULT_BASE,
                        help='Template fattura da usare come carta intestata')
    parser.add_argument('--out', default=DEFAULT_OUT,
                        help='Percorso del template giustificativo generato')
    args = parser.parse_args()

    doc = Document(args.base)
    strip_invoice_body(doc)
    append_giustificativo_body(doc)
    doc.save(args.out)
    print(f"Template salvato in {os.path.abspath(args.out)} (base: {args.base})")


if __name__ == '__main__':
    main()
